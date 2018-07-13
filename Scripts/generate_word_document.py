from docx import Document
from docx.shared import Inches
import csv
import os
import json

ruta_informes='/Users/dass/Tools/Results/Informes/'
ruta_permisos = '/Users/dass/Tools/Results/permisos/'
ruta_librerias = '/Users/dass/Tools/Results/librerias/processing_info/'
ruta_imagenes = '/Users/dass/Tools/Results/tables_static/images/'
ruta_dfs = '/Users/dass/Tools/Results/tables_static/df/'
ruta_jsons = '/Users/dass/Tools/Results/targets/'
ruta_codigo = '/Users/dass/Tools/Results/Codigo_fuente_apps/'
path_textfile = '/Users/dass/Tools/Results/fileslist.txt'

def table_from_csv(df, doc):
    with open(df, newline='') as f:
        csv_reader = csv.reader(f) 

        csv_headers = next(csv_reader)
        csv_cols = len(csv_headers)

        table = doc.add_table(rows=2, cols=csv_cols)
        hdr_cells = table.rows[0].cells
        table.style = 'TableGrid'
        table.autofit = True

        for i in range(csv_cols):
            hdr_cells[i].text = csv_headers[i]

        for row in csv_reader:
            row_cells = table.add_row().cells
            for i in range(csv_cols):
                row_cells[i].text = row[i]

    
def describir_libreria(libreria):
    if libreria == 'ADMOB':
        return 'AdMob, librería de anuncios que filtra ubicación a terceros. De hecho, Google ha eliminado unas cuantas aplicaciones que la utilizaban por problemas de seguridad.'
    elif libreria == 'ADWHIRL':
        return 'AdWhirl, librería centralizada de anuncios que recibe información de otros módulos de anuncios y filtra ubicación a terceros.'
    elif libreria == 'FLURRY' or libreria == 'FLURRY SDK':
        return 'Flurry, librería de anuncios que filtra ubicación a terceros.'
    elif libreria == 'INMOBI':
        return 'InMobi, librería de anuncios que filtra ubicación a terceros'
    elif libreria == 'MOPUB':
        return 'MoPub, librería de anuncios que filtra ubicación a terceros'
    elif libreria == 'APPLOVIN':
        return 'AppLovin, librería de anuncios que filtra ubicación a terceros.'
    elif libreria == 'ADGOJI':
        return 'Adgoji, librería de anuncios que filtra dirección MAC del usuario y sus gustos a terceros'
    elif libreria == 'ANDROID BEACON LIBRARY':
        return 'Android Beacon Library permite conectarse a balizas que pueden ser utilizadas como receptores y/o transmisores en algunas técnicas de fingerprinting.'
    elif libreria == 'ZXING':
        return 'ZXing, librería de lectura de códigos de barras convencionales y QR, que trata sin cuidado las URIs de los mismos, encargando su interpretación a terceros y pudiendo redirigir fácilmente esta información a la conveniencia del desarrollador de la app que usa la librería.'
    elif libreria == 'FLUZO':
        return 'Fluzo: Librería para la captación de audio y la construcción de una huella de usuario.'
    elif libreria == 'GRACENOTE':
        return 'GraceNote: Librería para la captación de audio y la construcción de una huella de usuario.'
    else:
        return libreria
    
def create_document(nombre_app, permisos, librerias, imagenes, dataframes, intellidroid):
    document = Document()

    # Cabecera
    document.add_heading('Informe de la aplicación {}'.format(nombre_app), 0)

    p = document.add_paragraph('Informe de la aplicación')
    p.add_run(' {}'.format(nombre_app)).bold = True
    p.add_run(', generado por el prototipo desarrollado.')

    # Permisos
    document.add_heading('Permisos', level=1)
    document.add_paragraph('Lista de permisos de la aplicacion y módulos a los que accede declarados en el Manifest', style='IntenseQuote')
    document.add_paragraph('Aqui se muestra la lista de permisos y módulos a los que puede acceder la aplicación según el permiso otorgado en el Manifest de la misma que pueden dar lugar a la construción de la huella digital. Esta lista sólo es un indicativo de que podría estarse realizando, no una garantía de que efectivamente esté utilizando dichos permisos para filtrar información privada.')

    for permiso in permisos:  
        document.add_paragraph(
            '{}'.format(permiso), style='ListBullet'
        )
    
    # Librerías
    document.add_heading('Librerías peligrosas', level=1)
    document.add_paragraph('Librerías de terceros que pueden ser utilizadas para filtrar información privada', style='IntenseQuote')
    document.add_paragraph('Se muestran las librerías de terceros que contiene la aplicación que forman parte de una lista negra que se ha creado para el prototipo. Estas librerías tienen filtraciones de información privada o pueden utilizarse para cometer filtraciones y/o construir la huella digital del usuario. De nuevo, puede que la aplicación en cuestión les esté dotando de otro uso diferente, pero puede que sea el uso fraudulento que comentamos.')

    for libreria in librerias: 
        libreria = describir_libreria(libreria)
        document.add_paragraph(
            '{}'.format(libreria), style='ListBullet'
        )
        
    # API de Android
    document.add_heading('Accesos a la API de Android por parte de la aplicación', level=1)
    document.add_paragraph('Funcionalidades de la API del sistema accedidas por la aplicación', style='IntenseQuote')
    document.add_paragraph('Información sobre las funciones de la interfaz de programación del sistema Android a las que accede la aplicación, con información detallada sobre qué clase o clases de la misma han accedido a la función remarcada. Mediante estos accesos la aplicación puede solicitar conexiones HTTP, crear sockets de red, acceder a la ubicación GPS, escribir y leer de la memoria de almacenamiento...')

    if dataframes[0] != '':
        table_from_csv(dataframes[0],document)
    if imagenes[0] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[0]))

    # Actividades e intents
    document.add_heading('Actividades e Intents', level=1)
    document.add_paragraph('Relación entre actividades que producen intents y la descripción de dichos intents', style='IntenseQuote')
    document.add_paragraph('Los "Intents" son mensajes asíncronos que permiten a los componentes de una aplicación Android solicitar una funcionalidad de otros componentes Android, permitiendo por tanto a la aplicación en cuestión interaccionar con otros componentes de la misma o con otras aplicaciones del dispositivo. Por ejemplo, acceder a la cámara de fotos y tomar una foto. Aquí se muestran las actividades de la aplicación que crean intents, y los datos del Intent. Estos datos pueden especificarse mediante su tipo MIME, prefijo URI (identificador de recurso), esquema URI (identificador del recurso en forma de esquema mediante su nombre y descripción) o una combinación de estos. Puede especificarse también el host de la URI, el puerto, el tipo MIME...')

    if dataframes[1] != '':
        table_from_csv(dataframes[1],document)
    if imagenes[1] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[1]))
    
    # Análisis del Manifest
    document.add_heading('Análisis del Manifest', level=1)
    document.add_paragraph('Análisis en detalle del Manifest de la aplicación y sus problemas de seguridad.', style='IntenseQuote')
    document.add_paragraph('En un apartado de arriba describimos los permisos y módulos a los que accede la aplicación descritos en el Manifest. Pero hay más información en el mismo, que puede ser peligrosa. En este apartado se muestran todos los problemas de seguridad detectados en el mismo, entre los que se incluirán algunos posibles problemas de privacidad, sobre todo relacionados con filtros puestos a intents para que pasen desapercibidos o receptores de red. Se incluye una descripción del problema y categorización según su peligro para el usuario.')

    if dataframes[2] != '':
        table_from_csv(dataframes[2],document)
    if imagenes[2] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[2]))
    
    # Análisis del código
    document.add_heading('Análisis del código', level=1)
    document.add_paragraph('Análisis estático del código com MobSF que detecta problemas de privacidad y seguridad declarados en el mismo.', style='IntenseQuote')
    document.add_paragraph('Resultado de analizar de forma estática (fuera del entorno de ejecución de la aplicación) el código puro de la misma. Esta parte puede revelar información enviada a logs, consultas a bases de datos, escritura a almacenamiento externo... Se indica asimismo la clase del código en que ocurre el problema detectado.')

    if dataframes[3] != '':
        table_from_csv(dataframes[3],document)
    if imagenes[3] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[3]))
    
    # Análisis de archivos peligrosos
    document.add_heading('Análisis de archivos', level=1)
    document.add_paragraph('Búsqueda de estructuras peligrosas en archivos de la aplicación.', style='IntenseQuote')
    document.add_paragraph('Este análisis es poco probable que produzca resultados, pero ahora se analizan los archivos directamente en lugar del código que contienen, en busca de estructuras peligrosas.')

    if dataframes[4] != '':
        table_from_csv(dataframes[4],document)
    if imagenes[4] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[4]))
    
    # URLs accedidas por la aplicación
    document.add_heading('URLs accedidas por la aplicación', level=1)
    document.add_paragraph('URLs a las que accede la aplicación', style='IntenseQuote')
    document.add_paragraph('La aplicación accede a las siguientes URLs, en concreto las clases especificadas a la derecha de la URL.')

    if dataframes[5] != '':
        table_from_csv(dataframes[5],document)
    if imagenes[5] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[5]))
    
    # E-mails enviados por la aplicación
    document.add_heading('Emails enviados por la aplicación', level=1)
    document.add_paragraph('Clases de la aplicación que pueden enviar emails', style='IntenseQuote')
    document.add_paragraph('Se especifican las clases del código que pueden realizar envío de emails a alguna dirección, con información sobre los mismos si procede.')

    if dataframes[6] != '':
        table_from_csv(dataframes[6],document)
    if imagenes[6] != '':
        document.add_paragraph('(Tiene esta información en formato de imagen en {}, por si desea una mejor visualización)'.format(imagenes[6]))
    
    # Targets de Intellidroid
    document.add_heading('Caminos y comportamientos que provocan filtraciones en código', level=1)
    document.add_paragraph('Información extraída mediante el análisis estático de Intellidroid', style='IntenseQuote')
    document.add_paragraph('Las aplicaciones que producen filtraciones sólo las llevan a cabo en partes determinadas del código, y si el análisis o el manejo del usuario no incurre en esa parte, no detectará la filtración. Esta parte del análisis busca vulnerabilidades en el código y muestra el camino de llamadas de la aplicación y comportamiento que permite incurrir en la vulnerabilidad que puede ocasionar la filtración de información. También nos ilustra sobre el tipo de filtración que se está llevando a cabo. Este apartado está pensado para que usted si lo desea siga el camino con las restricciones marcadas para detectar la posible filtración, sea visualmente o con una herramienta de análisis dinámico que lo automatice.')
    document.add_paragraph('AVISO: Si la aplicación tiene muchos caminos de vulnerabilidades, en este apartado pueden figurar muchas tablas que ocupen muchas páginas.')
    
    if intellidroid != {}:
        for i in intellidroid['callPaths']:
            document.add_paragraph('Camino de llamadas (call path) número {}:'.format(i), style='ListBullet').paragraph_format.left_indent = Inches(0.25)
            document.add_paragraph('Método que lo inicia: '+intellidroid['callPaths'][i]['startMethod']).paragraph_format.left_indent = Inches(0.25)
            document.add_paragraph ('Método objetivo: '+intellidroid['callPaths'][i]['targetMethod']).paragraph_format.left_indent = Inches(0.25)
            document.add_paragraph ('Cadena de eventos: ').paragraph_format.left_indent = Inches(0.25)
            for idx, val in enumerate(intellidroid['callPaths'][i]['eventChain']):
                columna_izda = []
                columna_dcha = []
                columna_izda.append ('Número de evento: ')
                columna_dcha.append ('Evento {}'.format(idx))
                columna_izda.append ('Inicio: ')
                columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['start'])
                columna_izda.append ('Objetivo: ')
                columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['target'])
                columna_izda.append ('Tipo de actividad: ')
                columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['type'])
                try:
                    columna_izda.append ('Archivo de restricciones necesarias para el camino: ')
                    columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['constraintsFile'])
                except:
                    columna_dcha.append ('No hay archivo de restricciones asociado.')
                    pass
                try:
                    columna_izda.append ('Componente: ')
                    columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['component'])
                except:
                    columna_dcha.append ('No hay componente.')
                    pass
                try:
                    columna_izda.append ('Variables involucradas: ')
                    columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['variables'])
                except:
                    columna_dcha.append ('No hay variables.')
                    pass
                try:
                    columna_izda.append ('Strings manejados: ')
                    columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['strings'])
                except:
                    columna_dcha.append ('No hay strings.')
                    pass
                try:
                    columna_izda.append ('Mapa de Strings manejados: ')
                    columna_dcha.append (intellidroid['callPaths'][i]['eventChain'][idx]['stringMap'])
                except:
                    columna_dcha.append ('No hay mapa de strings.')
                    pass
                table_in = document.add_table(rows = 1, cols = 2)
                table_in.autofit = True
                table_in.style = 'Table Grid'
                hdr_cells = table_in.rows[0].cells
                hdr_cells[0].text = 'Item'
                hdr_cells[1].text = 'Valor'
                for idx in range(0,len(columna_izda)):
                    row_cells = table_in.add_row().cells
                    row_cells[0].text = str(columna_izda[idx])
                    row_cells[1].text = str(columna_dcha[idx])
                document.add_paragraph ('\n')
    
    # Código fuente
    document.add_heading('Código fuente', level=1)
    document.add_paragraph('Código fuente de la aplicación', style='IntenseQuote')
    document.add_paragraph('El código fuente de la aplicación puede encontrarse en la carpeta {}'.format(ruta_codigo+nombre_app+'-jadx'))
    
    for target in intellidroid:  
        document.add_paragraph(
            '{}'.format(target), style='ListBullet'
        )
        
    # Consideraciones adicionales
    document.add_heading('Consideraciones adicionales', level=1)
    document.add_paragraph('', style='IntenseQuote')
    document.add_paragraph('En la carpeta de resultados se dispone del informe completo generado por MobSF de la aplicación, que incluye información sobre Malware de la aplicación además de la información de privacidad expresada aquí. Quizás esto también sea interesante para el usuario, pero no entra en el ámbito de este prototipo. Igualmente, puede consultarla en dicho informe.')
    document.add_paragraph('Puede que aunque no los muestre este análisis se estén cometiendo filtraciones adicionales, simplemente será indicativo de que la aplicación no se ha analizado correctamente o se escapa a las consideraciones que realiza este prototipo. Si el usuario desea más información o establer relaciones entre las filtraciones, puede utilizar una herramienta de análisis dinámico de las descritas en este TFM. Si elige utilizar TaintDroid, puede disparar los comportamientos deseados mediante los ficheros generados por Intellidroid en su carpeta del directorio de Resultados.')
    document.add_paragraph('Si desea un análisis rápido de la política de privacidad de la aplicación en cuestión, puede localizarla y utilizar sobre ella herramientas como Polisis (https://pribot.org/polisis) para obtener un resumen visual de la misma y comprobar si declara el uso de lo especificado en este informe.')

    document.add_page_break()

    document.save(ruta_informes+'Informe_{}.docx'.format(nombre_app))

    
if __name__ == '__main__':
    # We get the apps list
    with open(path_textfile) as f:
        lista_apps = f.readlines()
    lista_apps = [x.strip() for x in lista_apps] 
    # And all the files needed to show the result for each of them
    for app in lista_apps: 
        
        # Permission file
        archivo_permisos = ruta_permisos + 'permisos_'+app+'-jadx.txt'
        if os.path.exists(archivo_permisos):
            with open(ruta_permisos + 'permisos_'+app+'-jadx.txt') as p:
                permisos = p.readlines()
                permisos = [x.strip('\n') for x in permisos] 
        else:
            permisos = ['No se han podido obtener los permisos']
            
        # Libraries
        archivo_librerias = ruta_librerias + app +'.csv'
        librerias = []
        if os.path.exists(archivo_librerias):
            with open(archivo_librerias,'r') as arch:
                reader = csv.reader(arch,delimiter=',')
                for row in reader:
                    for element in row:
                        librerias.append(element)
        if len(librerias) == 0:
            librerias.append('No se han encontrado librerías de la lista negra en la aplicación')
            
        # Images from static analysis with MobSF
        
        imagenes = []
        for num in range(0,7):
            ruta_imagen = ruta_imagenes + app + '.apk_table_{}.png'.format(num)
            if os.path.exists(ruta_imagen):
                imagenes.append(ruta_imagen)
            else:
                imagenes.append('')
        
        # CSV Dataframes from processing MobSF analysis
        
        dataframes = []
        for num in range(0,7):
            ruta_df = ruta_dfs + app + '.apk_table_{}.csv'.format(num)
            if os.path.exists(ruta_df):
                dataframes.append(ruta_df)
            else:
                dataframes.append('')
        
        # And Intellidroid results
        json_file = ruta_jsons + app + '.json'
        if os.path.exists(json_file):
            json_intellidroid=open(json_file)
            intellidroid = json.load(json_intellidroid)
            json_intellidroid.close()
        else:
            intellidroid={}
        
        create_document(app, permisos, librerias, imagenes, dataframes, intellidroid)
        print('Informe generado para la aplicación '+app)
    print('Proceso terminado. Generados todos los informes')
        
        